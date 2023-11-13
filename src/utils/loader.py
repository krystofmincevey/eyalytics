import os
import re
import camelot
import tabula
import PyPDF2
import pdfplumber
import slate3k as slate
import docx

from typing import List, Dict
from tqdm.auto import tqdm
from pdfminer.high_level import extract_text
from docx2python import docx2python


PYPDF2_KEY = 'pypdf2'
SLATE_KEY = 'slate'
CAMELOT_KEY = 'camelot'
TABULA_KEY = 'tabula'
PLUMBER_KEY = 'plumber'
TIKA_KEY = 'tika'
MINER_KEY = 'miner'
REGEX_KEY = 'regex'
OCR_KEY = 'ocr'
DOXC_KEY = 'doxc'
DOXC2PYTHON_KEY = 'doxc2python'
FIGURE_THRESHOLD = 0.1
EPSILON = 1e-10
REPEAT_THRESHOLD = 4
MAX_CHAR_COUNT_FOR_FIGURE = 20
FIGURE_RELATED_CHARS = r"[0123456789.%-]"
COMMON_UNITS = ["kg", "m", "s", "h", "g", "cm", "mm", "l", "ml"]
DOCX_NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
# DICT KEYS
METADATA_KEY = "metadata"
COMP_KEY = "company"
REPORT_KEY = "report"
FIGURE_KEY = "potential_figure"
TABLE_KEY = "table"
TABLE_ROWS_KEY = "table"
TEXT_KEY = "text"


def read_txt(filename: str):
    with open(filename, 'r') as f:
        contents = f.read()
    return contents


def extract_footnotes_from_para(para, next_para=None):
    """Extract footnote references and actual footnotes from a paragraph."""
    footnotes = []

    footnote_refs = para._element.findall('.//w:footnoteReference', namespaces=DOCX_NAMESPACE)

    for ref in footnote_refs:
        footnote_id = ref.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
        footnote = para.part.footnotes_part.footnote_dict[footnote_id]
        footnotes.append(footnote.text)

    # Check in the next paragraph for footnotes if provided
    if next_para:
        next_footnote_refs = next_para._element.findall('.//w:footnoteReference', namespaces=DOCX_NAMESPACE)
        for ref in next_footnote_refs:
            footnote_id = ref.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
            footnote = next_para.part.footnotes_part.footnote_dict[footnote_id]
            footnotes.append(footnote.text)

    return footnotes


def process_footnotes(text, footnotes):
    """Process and embed footnotes into the text."""

    # Existing replacement for footnotes within brackets
    for idx, footnote in enumerate(footnotes, 1):
        text = re.sub(r"\[{}\]".format(idx), "[{}]".format(footnote), text)

    # New addition: replace footnotes appearing directly after words or at the end of sentences
    for idx, footnote in enumerate(footnotes, 1):
        # This regex will look for a number that doesn't have another number directly before
        # it (to differentiate from normal numbers within the text)
        pattern = r'(?<![0-9])' + str(idx) + r'(?![0-9])'
        replacement = "[{}]".format(footnote)
        text = re.sub(pattern, replacement, text)

    return text


def contains_unit(text):
    """Check if text contains a common unit following a number."""
    for unit in COMMON_UNITS:
        # Check for patterns like '123 kg', '123kg', '0.5 m', '0.5m', etc.
        if re.search(r'\d\s?' + re.escape(unit) + r'(?![a-zA-Z])', text):
            return True
    return False


def is_potential_figure_data(text):
    if text is None:
        return False

    text_count = len(text) - text.count(' ')
    figure_char_count = len(re.findall(FIGURE_RELATED_CHARS, text))
    char_count = text_count - figure_char_count

    # New: Check for units
    contains_common_units = contains_unit(text.lower())  # Convert text to lowercase for this check

    # New: Check for percentage patterns
    contains_percentage = "%" in text and any(char.isdigit() for char in text)

    if (text_count == 0) or text.endswith(('.', ':', ';', ',')) or (char_count > MAX_CHAR_COUNT_FOR_FIGURE):
        return False

    if (char_count == 0) or (figure_char_count / (
            char_count + EPSILON) > FIGURE_THRESHOLD) or contains_common_units or contains_percentage:
        return True

    return False


def repeated_artifact_check(line, artifact_dict):
    """Check if a line is a repeated artifact and update its count."""
    if line in artifact_dict:
        artifact_dict[line] += 1
        if artifact_dict[line] > REPEAT_THRESHOLD:
            return True  # It's a repeated artifact
    else:
        artifact_dict[line] = 1
    return False


def _is_empty(text):
    return len(text) == 0


class Loader(object):

    TEXT_LOADERS = {}
    TABLE_LOADERS = {}
    FIGURE_LOADERS = {}

    def __init__(self, file_path: str, dir_path: str):
        self._path = None  # Initialize _path to None
        self.path = os.path.join(dir_path, file_path)

    def load_tables(self, key: str):

        if key not in self.TABLE_LOADERS.keys():
            raise KeyError(
                f"The following table loaders: {[self.TABLE_LOADERS.keys()]} are supported. "
                f"Not {key}"
            )

        return self.TABLE_LOADERS[key](self.path)

    def load_text(self, key: str):

        if key not in self.TEXT_LOADERS.keys():
            raise KeyError(
                f"The following text loaders: {[self.TEXT_LOADERS.keys()]} are supported. "
                f"Not {key}"
            )

        return self.TEXT_LOADERS[key](self.path)

    def load_figure(self, key: str):

        if key not in self.FIGURE_LOADERS.keys():
            raise KeyError(
                f"The following figure loaders: {[self.FIGURE_LOADERS.keys()]} are supported. "
                f"Not {key}"
            )

        return self.FIGURE_LOADERS[key](self.path)


# TODO: Finalise, text, table, and figure loaders
class DOXCLoader(Loader):

    def __init__(self, file_path, dir_path='./data/doxc_db'):
        super().__init__(file_path=file_path, dir_path=dir_path)

    @property
    def path(self):
        return self._path

    @path.setter
    def path(self, new_path):
        if not os.path.exists(new_path):
            raise FileNotFoundError(f"The file {new_path} does not exist.")

        if not new_path.lower().endswith('.doxc'):
            raise ValueError("The file must be a DOXC.")

        if not os.path.isfile(new_path):
            raise ValueError("The path must point to a file, not a directory.")

        self._path = new_path

    def docx2python_text_loader(self) -> List[str]:
        """
        Extract text from a .doxc file using docx2python.
        :return: A list of strings, where each string represents a block of text.
        """
        # Extract the .doxc content
        docx_content = docx2python(self.path)
        text_content = []

        # docx2python represents the docx as a list of lists.
        # You navigate through these to find paragraphs.
        # Iterate through the body (ignores headers, footers)
        for docx_part in docx_content.body:
            for table in docx_part:
                for row in table:
                    for cell in row:
                        # strip() removes whitespace at the beginning and end of the text
                        processed_text = cell.text.strip()
                        if processed_text:  # ignore empty strings
                            text_content.append(processed_text)

        return text_content

    def doxc_text_loader(self) -> List[str]:
        """
        Extract text from a .doxc file.
        :return: A list of strings, where each string represents a block of text.
        """
        doc = docx.Document(self.path)
        text_content = []
        for paragraph in doc.paragraphs:
            processed_text = paragraph.text.strip()
            if not _is_empty(processed_text):
                text_content.append(processed_text)
        return text_content

    def doxc_fig_loader(self) -> List[Dict]:
        """
        Extract figures from a .doxc file.
        :return: A list of dicts, each containing the title and data of a figure.
        """
        doc = docx.Document(self.path)
        figures = []
        figure_data_group = {'title': None, 'data': []}
        previous_text = None

        for paragraph in doc.paragraphs:
            processed_text = paragraph.text.strip()
            if not _is_empty(processed_text):
                current_is_figure_data = is_potential_figure_data(processed_text)
                if current_is_figure_data:
                    if not figure_data_group['title']:
                        figure_data_group['title'] = previous_text
                    figure_data_group['data'].append(processed_text)
                else:
                    if figure_data_group['data']:
                        figures.append(figure_data_group)
                        figure_data_group = {'title': None, 'data': []}
                previous_text = processed_text
        return figures

    def doxc_table_loader(self) -> List[Dict]:
        """
        Extract tables from a .doxc file.
        :return: A list of dicts, each containing the title, column headers, and rows of a table.
        """
        doc = docx.Document(self.path)
        tables = []
        previous_text = None

        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []

            for row in table.rows[1:]:
                row_data = {headers[j]: cell.text.strip() for j, cell in enumerate(row.cells)}
                rows.append(row_data)

            tables.append({
                'title': previous_text,
                'col_headers': headers,
                TABLE_ROWS_KEY: rows
            })

            # Reset the previous_text if the title was just used
            previous_text = None

            # Update the previous_text with the last cell's text of the last row if it's not empty
            if rows and rows[-1]:
                last_cell_text = list(rows[-1].values())[-1]
                if last_cell_text.strip():
                    previous_text = last_cell_text

        return tables

    def doxc_load_all(
            self,
    ) -> dict:

        def _is_empty(text):
            if len(text) == 0:
                return True
            return False

        doc = docx.Document(self.path)

        result = {
            METADATA_KEY: {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
            },
            TEXT_KEY: [],
            TABLE_KEY: [],
            FIGURE_KEY: [],
        }

        figure_data_group = {'title': None, 'data': []}
        current_is_figure_data = False
        previous_text = None

        artifact_dict = {}

        for current_elem, next_elem in tqdm(zip(doc.element.body, doc.element.body[1:] + [None])):

            # Paragraph
            if current_elem.tag.endswith('p'):

                current_para = docx.text.paragraph.Paragraph(current_elem, None)
                next_para = docx.text.paragraph.Paragraph(next_elem, None)

                processed_text = current_para.text.strip()
                # Ignore empty lines or repeated lines
                if _is_empty(processed_text) or repeated_artifact_check(processed_text, artifact_dict):
                    continue
                try:
                    next_text = next_para.text
                except AttributeError:
                    next_text = None

                # Process footnotes
                footnotes = extract_footnotes_from_para(current_para, next_para)
                processed_text = process_footnotes(processed_text, footnotes)

                # Identify if the current line is potential figure data
                previous_was_figure_data = current_is_figure_data  # Move the window forward
                current_is_figure_data = is_potential_figure_data(processed_text)
                next_is_figure_data = is_potential_figure_data(next_text)

                if current_is_figure_data:

                    # If previous line was also figure data, they belong to the same figure
                    if previous_was_figure_data:
                        figure_data_group['data'].append(processed_text)
                    else:
                        # If a new figure starts, save the previous figure (if there was any)
                        if figure_data_group['data']:
                            result[FIGURE_KEY].append(figure_data_group)
                            figure_data_group = {'title': None, 'data': []}

                        # Assign the previous line as the title for the current figure
                        figure_data_group['title'] = previous_text
                        figure_data_group['data'].append(processed_text)

                elif not next_is_figure_data:  # neither next or current text is figure
                    # Not a figure, add to text
                    result[TEXT_KEY].append(processed_text)
                else:  # next text is figure, meaning that current text will be stored as title.
                    pass

                # Handles case when text potentially interrupts figure.
                # Text is again stored in figure_data_group['data'].
                if previous_was_figure_data and next_is_figure_data:
                    current_is_figure_data = True

                # only change previous text when current para is not empty or repeated string.
                previous_text = processed_text

            # Table
            elif current_elem.tag.endswith('tbl'):
                table_index = [tbl._element for tbl in doc.tables].index(current_elem)
                table = doc.tables[table_index]

                headers = [cell.text.strip() for cell in table.rows[0].cells]

                rows = []
                for row in table.rows[1:]:
                    row_data = {headers[j]: cell.text.strip() for j, cell in enumerate(row.cells)}
                    rows.append(row_data)

                result[TABLE_KEY].append({
                    'title': previous_text,
                    'col_headers': headers,
                    TABLE_ROWS_KEY: rows
                })
            else:
                print(f"Ignoring {current_elem.tag}.")

        return result

    TEXT_LOADERS = {
        DOXC_KEY: doxc_text_loader,
        DOXC2PYTHON_KEY: docx2python_text_loader,
    }
    TABLE_LOADERS = {
        DOXC_KEY: doxc_table_loader,
    }
    FIGURE_LOADERS = {
        DOXC_KEY: doxc_fig_loader,
    }


class PDFLoader(Loader):

    def __init__(self, file_path, dir_path='./data/pdf_db'):
        super().__init__(file_path=file_path, dir_path=dir_path)

    @property
    def path(self):
        return self._path

    @path.setter
    def path(self, new_path):
        if not os.path.exists(new_path):
            raise FileNotFoundError(f"The file {new_path} does not exist.")

        if not new_path.lower().endswith('.pdf'):
            raise ValueError("The file must be a PDF.")

        if not os.path.isfile(new_path):
            raise ValueError("The path must point to a file, not a directory.")

        self._path = new_path

    @staticmethod
    def pypdf2_text_loader(file_path: str) -> List[str]:
        """
        Extract text from a PDF file using PyPDF2.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents the text content of a page.
        """
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            text_content = []
            for page_num in range(reader.getNumPages()):
                page = reader.getPage(page_num)
                text_content.append(page.extract_text())
        return text_content

    @staticmethod
    def pdfminer_text_loader(file_path: str) -> List[str]:
        """
        Extract text from a PDF file using pdfminer.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents a line of text.
        """
        text_content = extract_text(file_path)
        return text_content.split('\n')  # Splitting by newline to get a list of lines

    @staticmethod
    def slate_text_loader(file_path) -> List[str]:
        """
        Extract text from a PDF file using slate.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents the text content of a page.
        """
        with open(file_path, 'rb') as f:
            doc = slate.PDF(f)
        text_content = [page for page in doc]
        # You would need additional logic here to separate tables from text
        return text_content

    @staticmethod
    def regex_table_loader(file_path) -> List[str]:
        """
        Extract tables from a PDF file using regular expressions.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents a detected table row.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()

        # This is a very simplistic example and may need to be adjusted
        # to suit the actual format of your tables.
        tables = re.findall(r'(\d+ \d+ \d+)', text_content)
        return tables

    @staticmethod
    def pdfplumber_text_loader(file_path) -> List[str]:
        """
        Extract text from a PDF file using pdfplumber.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents the text content of a page.
        """
        with pdfplumber.open(file_path) as pdf:
            text_content = [page.extract_text() for page in pdf.pages]
        return text_content

    @staticmethod
    def pdfplumber_table_loader(file_path) -> List[str]:
        """
        Extract tables from a PDF file using pdfplumber.

        :param file_path: The path to the PDF file.
        :return: A list of tables, where each table is represented as a list of rows,
                 and each row is a list of strings.
        """
        with pdfplumber.open(file_path) as pdf:
            tables = []
            for page in pdf.pages:
                page_tables = page.extract_tables()
                tables.extend(page_tables)

        # Convert the tables to list of strings
        result = []
        for table in tables:
            for row in table:
                result.append(' '.join(row))
        return result

    @staticmethod
    def pypdf2_table_loader(file_path) -> List[str]:
        """
        Extract text and tables from a PDF file using an enhanced version of PyPDF2.

        :param file_path: The path to the PDF file.
        :return: A list of pages, where each page is represented as a list of strings.
                 Strings represent either lines of text or detected table rows.
        """
        def enhanced_text_analysis(text):
            # Implement custom logic to analyze and extract tables or structured data
            # This could involve regular expressions, heuristics, or machine learning models

            # For example, a simple heuristic might be to look for lines of text that are
            # formatted like table rows, perhaps with data separated by spaces or other delimiters
            lines = text.split('\n')
            tables = []
            current_table = []
            for line in lines:
                if is_table_row(line):  # You need to define is_table_row function
                    current_table.append(line)
                elif current_table:
                    tables.append(current_table)
                    current_table = []
            if current_table:
                tables.append(current_table)

            return tables

        def is_table_row(line):
            # Define your own logic to determine if a line of text represents a table row
            # This is a simplistic example and might need to be significantly expanded
            # depending on your PDFs' content and structure
            if len(line.split()) > 1:  # If the line has multiple words or numbers, consider it a table row
                return True
            return False

        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfFileReader(f)
            text_content = []
            for page_num in range(reader.getNumPages()):
                page = reader.getPage(page_num)
                text_content.append(enhanced_text_analysis(page.extract_text()))

        # Convert the tables to list of strings
        result = []
        for page_content in text_content:
            for item in page_content:
                if isinstance(item, list):  # If the item is a table (list of lists)
                    for row in item:
                        result.append(' '.join(row))
                else:
                    result.append(item)
        return result

    @staticmethod
    def camelot_table_loader(file_path: str) -> List[str]:
        """
        Extract tables from a PDF file using Camelot.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents a table row.
        """
        tables = camelot.read_pdf(file_path, pages='all', flavor='stream')
        result = []
        for table in tables:
            df = table.df
            for index, row in df.iterrows():
                result.append(' '.join(row))
        return result

    @staticmethod
    def tabula_table_loader(file_path: str) -> List[str]:
        """
        Extract tables from a PDF file using Tabula.

        :param file_path: The path to the PDF file.
        :return: A list of strings, where each string represents a table row.
        """
        dfs = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
        result = []
        for df in dfs:
            for index, row in df.iterrows():
                result.append(' '.join(row.astype(str)))
        return result

    @staticmethod
    def ocr_loader(file_path) -> List[str]:
        # TODO: Implement OCR solution using (for instance) pytesseract and cv2
        raise NotImplementedError

    TEXT_LOADERS = {
        PYPDF2_KEY: pypdf2_text_loader,
        SLATE_KEY: slate_text_loader,
        PLUMBER_KEY: pdfplumber_text_loader,
        MINER_KEY: pdfminer_text_loader,
    }

    # TODO: Consider LayoutLM (by Microsoft)
    TABLE_LOADERS = {
        CAMELOT_KEY: camelot_table_loader,
        TABULA_KEY: tabula_table_loader,
        PYPDF2_KEY: pypdf2_table_loader,
        PLUMBER_KEY: pdfplumber_table_loader,
        REGEX_KEY: regex_table_loader,
        OCR_KEY: ocr_loader,
    }
